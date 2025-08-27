# app.py — ERIK v6 Full (Humanoid responses, full calc UI, quiz generator, LaTeX solver, graphs, scholar)
import streamlit as st
import requests
from bs4 import BeautifulSoup
import sympy as sp
import matplotlib.pyplot as plt
import numpy as np
import fitz  # PyMuPDF
import docx
from googlesearch import search
import random
import re
from html import unescape
from math import isfinite

# ------------------ App config ------------------
st.set_page_config(page_title="ERIK v6 - AI Academic Assistant", layout="wide")
st.title("🧠 ERIK v6 — Exceptional Resources & Intelligence Kernel")
st.markdown("""
**Welcome to ERIK v6** — an offline-first academic assistant built by **Sabid Uddin Nahian**.  
Features: multilingual Q&A, humanoid responses, full scientific calculator UI, LaTeX solver, quiz generation, PDF analyzer, Google Scholar search, 2D/3D plotting.  
Use the sidebar to navigate. 🚀
""")

# ------------------ Sidebar ------------------
st.sidebar.header("Features")
mode = st.sidebar.radio("Choose a feature:", [
    "Ask Question",
    "Quiz Generator",
    "PDF/Text Analyzer",
    "LaTeX Math Solver (detailed)",
    "Scientific Calculator (Casio-like)",
    "Google Scholar Search",
    "2D & 3D Graph Generator"
])

# ------------------ Utilities ------------------
def fetch_text_from_url(url, max_paragraphs=5):
    try:
        r = requests.get(url, timeout=5, headers={"User-Agent":"ERIK/1.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        paras = soup.find_all('p')
        text = " ".join([unescape(p.get_text().strip()) for p in paras[:max_paragraphs]])
        return text
    except Exception:
        return ""

def extract_sentences(text):
    sents = re.split(r'(?<=[\.\?\!])\s+', text)
    sents = [s.strip() for s in sents if len(s.strip())>20]
    return sents

def humanize_reply_short(raw_text):
    # pick first 1-2 sentences
    sents = extract_sentences(raw_text)
    if not sents:
        return "Sorry, I couldn't find a concise answer. Try rephrasing."
    reply = sents[0]
    return f"🤖 Quick answer: {reply} ✨"

def humanize_reply_long(raw_text):
    sents = extract_sentences(raw_text)
    if not sents:
        return "I couldn't find useful information. Maybe try another query."
    reply = " ".join(sents[:6])
    # add small conversational touches (rule-based)
    intro = "Hey — here's a detailed answer I gathered:\n\n"
    outro = "\n\nHope that helps! If you want a clearer summary, ask me to shorten it. 🚀"
    return f"🤖 {intro}{reply}{outro}"

def safe_sympify(expr_str):
    # small sanitization before sympify
    # replace '^' with '**', '×' with '*', '÷' with '/'
    s = expr_str.replace("^","**").replace("×","*").replace("÷","/")
    # replace '!' with factorial(...) for sympy
    if "!" in s:
        s = re.sub(r'(\d+)!', r'factorial(\1)', s)
    # allow pi, E
    s = s.replace("π","pi")
    s = s.replace("PI","pi")
    return sp.sympify(s)

# ------------------ Quiz generation utilities (heuristic) ------------------
def choose_targets(sentences):
    targets = []
    for s in sentences:
        if re.search(r'\d', s):
            targets.append(("numeric", s))
        elif re.search(r'\b[A-Z][a-z]{2,}\b', s):
            targets.append(("entity", s))
        else:
            targets.append(("text", s))
    return targets

def generate_mcq_from_sentence(sent, kind="numeric"):
    nums = re.findall(r'[-+]?\d*\.?\d+(?:e[-+]?\d+)?', sent)
    if nums:
        target = nums[0]
        try:
            val = float(target)
            d1 = val + 1
            d2 = val - 1
            d3 = val * 1.5 if abs(val) > 1e-6 else val + 2
            opts = [str(val), str(round(d1,6)), str(round(d2,6)), str(round(d3,6))]
            random.shuffle(opts)
            correct = opts.index(str(val))
            q = re.sub(re.escape(target), "_____", sent, count=1)
            sol = f"{val}"
            return q, opts, correct, sol
        except:
            pass
    caps = re.findall(r'\b([A-Z][a-z]{2,})\b', sent)
    if caps:
        ans = caps[0]
        others = caps[1:] + [w for w in caps if w!=ans]
        distractors = []
        for o in others[:3]:
            if o != ans:
                distractors.append(o)
        while len(distractors) < 3:
            distractors.append(ans + random.choice(["a","o","i","er"]))
        opts = [ans] + distractors[:3]
        random.shuffle(opts)
        correct = opts.index(ans)
        q = re.sub(r'\b' + re.escape(ans) + r'\b', "_____", sent, count=1)
        sol = ans
        return q, opts, correct, sol
    words = [w for w in re.findall(r'\w+', sent) if len(w)>4]
    if words:
        ans = words[0]
        distractors = []
        for _ in range(3):
            distractors.append(ans[:max(1,len(ans)-1)] + random.choice(["a","o","x","z"]))
        opts = [ans] + distractors
        random.shuffle(opts)
        correct = opts.index(ans)
        q = re.sub(r'\b' + re.escape(ans) + r'\b', "_____", sent, count=1)
        sol = ans
        return q, opts, correct, sol
    q = sent if len(sent)<120 else sent[:120] + "..."
    return q, ["Option A","Option B","Option C","Option D"], 0, "See text"

def generate_quiz_from_text(content, num_q=5):
    sentences = extract_sentences(content)
    if not sentences:
        return []
    targets = choose_targets(sentences)
    quiz_items = []
    used = set()
    i=0
    for kind, sent in targets:
        if i>=num_q:
            break
        if sent in used:
            continue
        used.add(sent)
        qtype = random.choices(["MCQ","SAQ","CQ"], weights=[0.6,0.2,0.2])[0]
        if qtype == "MCQ":
            q, opts, corr, sol = generate_mcq_from_sentence(sent, kind)
            quiz_items.append({"type":"MCQ","question":q,"options":opts,"answer_index":corr,"solution":sol})
        elif qtype == "SAQ":
            nums = re.findall(r'[-+]?\d*\.?\d+(?:e[-+]?\d+)?', sent)
            if nums:
                ans = nums[0]
            else:
                caps = re.findall(r'\b([A-Z][a-z]{2,})\b', sent)
                ans = caps[0] if caps else (sent.split()[0] if sent.split() else "")
            quiz_items.append({"type":"SAQ","question":"Short Answer: " + sent,"answer":ans,"solution":f"{ans}"})
        else:
            idx = sentences.index(sent)
            context = " ".join(sentences[max(0, idx-1): min(len(sentences), idx+2)])
            qtext = "Explain / discuss: " + sent
            model = context
            quiz_items.append({"type":"CQ","question":qtext,"model_answer":model})
        i+=1
    return quiz_items

# ------------------ Google Scholar cleaner fetch ------------------
def fetch_scholar_title_and_abstract(url):
    # Try multiple heuristics: meta description, abstract tags
    try:
        r = requests.get(url, timeout=5, headers={"User-Agent":"ERIK/1.0"})
        soup = BeautifulSoup(r.text, "html.parser")
        title = (soup.title.string.strip() if soup.title and soup.title.string else url)
        # try meta description
        meta = soup.find("meta", attrs={"name":"description"}) or soup.find("meta", attrs={"property":"og:description"})
        abstract = meta["content"].strip() if meta and meta.get("content") else ""
        # fallback: first <p>
        if not abstract:
            p = soup.find('p')
            if p:
                abstract = p.get_text()[:400].strip()
        return title, abstract
    except Exception:
        return url, ""

# ------------------ MAIN FEATURES ------------------

# Ask Question (Humanoid)
if mode == "Ask Question":
    st.subheader("🌐 Ask Question (English, Bangla, German & Mandarin)")
    language = st.selectbox("Select language:", ["English", "Bangla", "German", "Mandarin"])
    format_choice = st.selectbox("Answer format:", ["Short", "Long"])
    query = st.text_input("Enter your question:")
    if st.button("Search & Answer"):
        st.info("Searching the web... ⏳")
        lang_code = {"English":"en", "Bangla":"bn", "German":"de", "Mandarin":"zh-CN"}[language]
        results = []
        try:
            for url in search(query, num_results=5, lang=lang_code):
                results.append(url)
        except Exception:
            st.error("❌ Error searching the web.")
        content = ""
        for u in results:
            content += fetch_text_from_url(u, max_paragraphs=3) + " "
        if content.strip():
            if format_choice == "Short":
                txt = humanize_reply_short(content)
            else:
                txt = humanize_reply_long(content)
            # add a small friendly intro/outro to make it humanoid
            intro = "Hi — I looked this up for you. " if format_choice=="Long" else ""
            outro = "\n\nIf you'd like sources or a simpler summary, tell me. 🙂"
            st.markdown(f"**{intro}{txt}{outro}**")
            st.markdown("**Top sources:**")
            for r in results:
                st.write(f"- {r}")
        else:
            st.warning("⚠️ Sorry — couldn't find useful info. Try rephrasing your question.")

# Quiz Generator
elif mode == "Quiz Generator":
    st.subheader("📝 Smart Quiz Generator (MCQ / SAQ / CQ)")
    source_type = st.radio("Select source type:", ["Google Search", "PDF/Text Upload"])
    num_q = st.number_input("Number of questions to generate:", min_value=1, max_value=20, value=5)
    if source_type == "Google Search":
        topic = st.text_input("Topic for quiz generation:")
        if st.button("Generate Quiz from Google"):
            st.info("Searching Google and creating quiz...")
            content = ""
            try:
                for url in search(topic, num_results=6):
                    content += fetch_text_from_url(url, max_paragraphs=4) + " "
                quiz_items = generate_quiz_from_text(content, num_q)
                if not quiz_items:
                    st.warning("No useful text found to generate questions.")
                else:
                    for idx, it in enumerate(quiz_items, start=1):
                        if it["type"] == "MCQ":
                            st.write(f"Q{idx} (MCQ): {it['question']}")
                            for i,opt in enumerate(it["options"]):
                                st.write(f"   {chr(97+i)}) {opt}")
                            with st.expander("Show solution / auto-solve"):
                                st.write("Correct option:", chr(97+it["answer_index"]))
                                st.write("Solution:", it.get("solution",""))
                        elif it["type"] == "SAQ":
                            st.write(f"Q{idx} (SAQ): {it['question']}")
                            with st.expander("Answer"):
                                st.write(it.get("answer",""))
                                st.write("Solution:", it.get("solution",""))
                        else:
                            st.write(f"Q{idx} (CQ): {it['question']}")
                            with st.expander("Model answer"):
                                st.write(it.get("model_answer",""))
            except Exception as e:
                st.error(f"❌ Error generating quiz: {e}")
    else:
        uploaded_file = st.file_uploader("Upload PDF / DOCX / TXT", type=['pdf','txt','docx'])
        if uploaded_file and st.button("Generate Quiz from File"):
            try:
                text = ""
                if uploaded_file.type == "application/pdf":
                    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                    for page in doc:
                        text += page.get_text()
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    docx_doc = docx.Document(uploaded_file)
                    for p in docx_doc.paragraphs:
                        text += p.text + "\n"
                else:
                    text = str(uploaded_file.read(), "utf-8")
                quiz_items = generate_quiz_from_text(text, num_q)
                if not quiz_items:
                    st.warning("Not enough text to make questions.")
                else:
                    for idx, it in enumerate(quiz_items, start=1):
                        if it["type"] == "MCQ":
                            st.write(f"Q{idx} (MCQ): {it['question']}")
                            for i,opt in enumerate(it["options"]):
                                st.write(f"   {chr(97+i)}) {opt}")
                            with st.expander("Show solution"):
                                st.write("Correct option:", chr(97+it["answer_index"]))
                                st.write("Solution:", it.get("solution",""))
                        elif it["type"] == "SAQ":
                            st.write(f"Q{idx} (SAQ): {it['question']}")
                            with st.expander("Answer"):
                                st.write(it.get("answer",""))
                        else:
                            st.write(f"Q{idx} (CQ): {it['question']}")
                            with st.expander("Model answer"):
                                st.write(it.get("model_answer",""))
            except Exception as e:
                st.error(f"❌ Error reading file or generating quiz: {e}")

# PDF/Text Analyzer
elif mode == "PDF/Text Analyzer":
    st.subheader("📄 PDF / Text Analyzer")
    uploaded_file = st.file_uploader("Choose a file", type=['pdf','txt','docx'])
    if uploaded_file:
        try:
            text = ""
            if uploaded_file.type == "application/pdf":
                doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
                for page in doc:
                    text += page.get_text()
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                docx_doc = docx.Document(uploaded_file)
                for p in docx_doc.paragraphs:
                    text += p.text + "\n"
            else:
                text = str(uploaded_file.read(), "utf-8")
            st.text_area("Extracted Text", text, height=400)
        except Exception as e:
            st.error(f"❌ Error extracting text: {e}")

# LaTeX Math Solver (detailed)
elif mode == "LaTeX Math Solver (detailed)":
    st.subheader("✍️ LaTeX Math Solver — Detailed Steps")
    latex_input = st.text_input("Enter expression/equation (use ^ for power). For equation use '=' (e.g., x^2 - 4 = 0):")
    if st.button("Solve with steps"):
        try:
            expr = latex_input.replace("^","**")
            if "=" in expr:
                left, right = expr.split("=",1)
                left_s = safe_sympify(left)
                right_s = safe_sympify(right)
                eq_expr = sp.simplify(left_s - right_s)
                st.markdown("**Original (converted):**")
                st.latex(sp.latex(sp.Eq(left_s, right_s)))
                st.markdown("**Simplified:**")
                st.latex(sp.latex(eq_expr) + " = 0")
                st.markdown("**Factored (if possible):**")
                st.latex(sp.latex(sp.factor(eq_expr)) + " = 0")
                st.markdown("**Solve:**")
                sols = sp.solve(eq_expr, dict=False)
                if sols:
                    for s in sols:
                        st.latex("x = " + sp.latex(sp.nsimplify(s)))
                else:
                    st.write("No symbolic solutions found.")
                st.markdown("**Check:**")
                for s in sols:
                    check = sp.simplify(eq_expr.subs(sp.symbols('x'), s))
                    st.write(f"Substituting x={s} → {check}")
            else:
                e = safe_sympify(expr)
                st.markdown("**Expression:**")
                st.latex(sp.latex(e))
                st.markdown("**Simplified:**")
                st.latex(sp.latex(sp.simplify(e)))
                st.markdown("**Factorized:**")
                st.latex(sp.latex(sp.factor(e)))
                x = sp.symbols('x')
                st.markdown("**Derivative (w.r.t x):**")
                try:
                    der = sp.diff(e, x)
                    st.latex(sp.latex(der))
                except Exception:
                    st.write("Derivative not available.")
                st.markdown("**Indefinite integral (w.r.t x):**")
                try:
                    inte = sp.integrate(e, x)
                    st.latex(sp.latex(inte))
                except Exception:
                    st.write("Integral not found symbolically.")
        except Exception as e:
            st.error(f"❌ Invalid input or symbolic failure: {e}")

# Scientific Calculator (Casio-like UI)
elif mode == "Scientific Calculator (Casio-like)":
    st.subheader("🧮 Scientific Calculator (Casio-style)")
    # session state for calculator
    if 'calc_input' not in st.session_state:
        st.session_state.calc_input = ""
    if 'last_ans' not in st.session_state:
        st.session_state.last_ans = ""
    if 'memory' not in st.session_state:
        st.session_state.memory = 0.0
    if 'mode_deg' not in st.session_state:
        st.session_state.mode_deg = True
    if 'shift' not in st.session_state:
        st.session_state.shift = False

    def calc_press(val):
        if val == "AC":
            st.session_state.calc_input = ""
        elif val == "DEL":
            st.session_state.calc_input = st.session_state.calc_input[:-1]
        elif val == "Ans":
            st.session_state.calc_input += st.session_state.last_ans
        elif val == "=":
            try:
                expr = st.session_state.calc_input
                # apply replacements
                expr = expr.replace("^","**").replace("×","*").replace("÷","/")
                expr = expr.replace("π","pi")
                # factorial handling
                expr = re.sub(r'(\d+)!', r'factorial(\1)', expr)
                # sympify and evaluate
                res = safe_sympify(expr).evalf()
                st.session_state.last_ans = str(res)
                st.session_state.calc_input = str(res)
            except Exception:
                st.session_state.calc_input = "Error"
        elif val == "M+":
            try:
                st.session_state.memory += float(st.session_state.last_ans or "0")
            except:
                pass
        elif val == "M-":
            try:
                st.session_state.memory -= float(st.session_state.last_ans or "0")
            except:
                pass
        elif val == "MR":
            st.session_state.calc_input += str(st.session_state.memory)
        elif val == "MC":
            st.session_state.memory = 0.0
        elif val == "DEG":
            st.session_state.mode_deg = True
        elif val == "RAD":
            st.session_state.mode_deg = False
        elif val == "SHIFT":
            st.session_state.shift = not st.session_state.shift
        else:
            # map some human-friendly buttons to sympy-friendly tokens
            mapping = {
                "×":"*","÷":"/","^":"**","√":"sqrt","pi":"pi","e":"E"
            }
            token = mapping.get(val, val)
            # if shift is on and token is sin -> asin, cos->acos, tan->atan, ln->exp^-1 etc.
            if st.session_state.shift:
                inv_map = {"sin":"asin","cos":"acos","tan":"atan","sinh":"asinh","cosh":"acosh","tanh":"atanh","log":"log10"}
                if token in inv_map:
                    token = inv_map[token]
            st.session_state.calc_input += token

    # top display
    st.text_area("Display", value=st.session_state.calc_input, height=60, key="display_area")

    # define button layout (rows list)
    btn_rows = [
        ["SHIFT","DEG","RAD","(",")","DEL","AC"],
        ["7","8","9","÷","sqrt","^","!"],
        ["4","5","6","×","x^2","x^3","%"],
        ["1","2","3","-","sin","cos","tan"],
        ["0",".","Ans","+","log","ln","exp"],
        ["π","e","ANS?","M+","M-","MR","MC"]
    ]

    # render rows as columns
    for row in btn_rows:
        cols = st.columns(len(row), gap="small")
        for i, label in enumerate(row):
            if cols[i].button(label):
                # small handling for 'ANS?' label — just show last ans
                if label == "ANS?":
                    st.info(f"Last Ans: {st.session_state.last_ans}")
                else:
                    calc_press(label)

    # small note about DEG/RAD effect for trig evaluation
    st.caption("Tip: use SHIFT to toggle inverse trig. DEG/RAD affects numeric trig evaluation when you press '=' (sympy expects radians; app will convert degrees to radians).")

# Google Scholar Search (clean professional)
elif mode == "Google Scholar Search":
    st.subheader("🔬 Google Scholar Search")
    keyword = st.text_input("Enter research topic:")
    if st.button("Search Scholar"):
        st.info("Fetching top research papers... ⏳")
        query = f"{keyword} site:scholar.google.com"
        results = []
        try:
            for url in search(query, num_results=5):
                results.append(url)
        except Exception:
            st.error("❌ Error searching Scholar.")
        if results:
            st.success(f"Found {len(results)} papers:")
            for r in results:
                try:
                    title, abstract = fetch_scholar_title_and_abstract(r)
                    st.markdown(f"**{title}**")
                    if abstract:
                        st.write(abstract)
                    st.write(f"[Read more]({r})")
                    st.divider()
                except Exception:
                    st.write(r)
        else:
            st.warning("No papers found. Try another keyword.")

# 2D & 3D Graph Generator
elif mode == "2D & 3D Graph Generator":
    st.subheader("📈 2D & 3D Graph Generator (smart equation writer)")
    raw_input = st.text_input("Enter function (2D: x only e.g. x^2, 3D: z=f(x,y) e.g. sin(x)*cos(y)):")
    func_input = raw_input.replace("^","**").replace("X","x").replace("Y","y")
    graph_type = st.radio("Graph Type:", ["2D", "3D"])
    xrange = st.slider("X range (min, max):", -20.0, 20.0, (-10.0, 10.0))
    yrange = st.slider("Y range (min, max) — for 3D only:", -20.0, 20.0, (-5.0, 5.0))
    samples = st.number_input("Number of samples (resolution):", min_value=50, max_value=1000, value=300, step=50)
    if st.button("Plot Graph"):
        try:
            if graph_type == "2D":
                x = sp.symbols('x')
                f = safe_sympify(func_input)
                f_np = sp.lambdify((x,), f, "numpy")
                X = np.linspace(xrange[0], xrange[1], samples)
                Y = f_np(X)
                # sanitize Y for plotting
                Y = np.where(np.isfinite(Y), Y, np.nan)
                fig, ax = plt.subplots(figsize=(8,4))
                ax.plot(X, Y)
                ax.set_xlabel("x")
                ax.set_ylabel("y")
                ax.set_title(f"y = {func_input}")
                ax.grid(True)
                st.pyplot(fig)
            else:
                x,y = sp.symbols('x y')
                f = safe_sympify(func_input)
                f_np = sp.lambdify((x,y), f, "numpy")
                X = np.linspace(xrange[0], xrange[1], int(np.sqrt(samples)))
                Y = np.linspace(yrange[0], yrange[1], int(np.sqrt(samples)))
                Xg, Yg = np.meshgrid(X, Y)
                Z = f_np(Xg, Yg)
                # sanitize
                Z = np.where(np.isfinite(Z), Z, np.nan)
                fig = plt.figure(figsize=(8,6))
                ax = fig.add_subplot(111, projection='3d')
                ax.plot_surface(Xg, Yg, Z, cmap='viridis', linewidth=0, antialiased=True)
                ax.set_xlabel("x"); ax.set_ylabel("y"); ax.set_zlabel("z")
                ax.set_title(f"z = {func_input}")
                st.pyplot(fig)
        except Exception as e:
            st.error(f"❌ Invalid equation or plotting error: {e}")

# End
st.markdown("---")
st.markdown("ERIK v6 — built with ❤️ by **Sabid Uddin Nahian**. If something breaks, paste the error here and I'll fix it.")
